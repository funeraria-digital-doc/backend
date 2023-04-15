import io
import os
from django.http import HttpResponse
from django.shortcuts import render
from rest_framework.response import Response
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from django.utils.decorators import method_decorator
from rest_framework.permissions import IsAuthenticated
from docxtpl import DocxTemplate
from accounts.models import User
from template_logic.serealizers import SEND_TYPE_CHOICES
from template_logic.serealizers import DB_COLLECTION_CHOICES
from template_logic.validation_helper import *
from template_logic.variables import get_var_types
from groups.models import Group
from template_logic.serealizers import UploadSerializer, EditUploadSerializer
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from rest_framework import generics
from template_logic.models import TemplateLogic
from records.models import Record
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
from rest_framework.parsers import MultiPartParser, JSONParser
from rest_framework.decorators import parser_classes
import json
from rest_framework.authtoken.models import Token
from docx import Document

@swagger_auto_schema(
    method='get',
    operation_description="List all templates"
) 
@api_view(['GET'])
def list_templates(request):
    templates = TemplateLogic.objects.all().values()
    if(templates.count() > 0):
        return Response({'data' : templates}, status=status.HTTP_200_OK)
    return Response({'error' : 'No templates available'}, status=status.HTTP_400_BAD_REQUEST)

@swagger_auto_schema(
    method='get',
    operation_description="Get all variables from a template"
) 
@api_view(['GET'])
#@permission_classes([IsAuthenticated])
def get_variables(request, *args, **kwargs):
    template = TemplateLogic.objects.filter(id=kwargs.get('pk')).values().first()
    if template is None:
        return Response({"error" : "Template does not exist!"}, status = status.HTTP_404_NOT_FOUND)
    variables = get_doc_variables(template)['vars']
    return Response(variables, status = status.HTTP_200_OK)

@swagger_auto_schema(
    method='post',
    request_body=UploadSerializer,
    operation_description="Upload a document template"
)    
@api_view(['POST'])
@parser_classes([MultiPartParser])
def upload(request):
    dup = hasDuplicates(request.POST.get('validations'))
    if dup:
        return dup
    data = {
        'title' : request.POST.get('title'),
        'group' : request.POST.get('group'),
        'file' : request.FILES.get('file'),
        'validations' : json.loads(request.POST.get('validations')),
        'send_type' : json.loads(request.POST.get('send_type')),
        'send_email_to' : json.loads(request.POST.get('send_email_to')) if request.POST.get('send_email_to') else [],
        'send_email_to_cc' : json.loads(request.POST.get('send_email_to_cc')) if request.POST.get('send_email_to_cc') else [],
        'send_email_to_bcc' : json.loads(request.POST.get('send_email_to_bcc')) if request.POST.get('send_email_to_bcc') else [],
    }
    form = UploadSerializer(data = data)
    if form.is_valid():
        try:
            if data['file']:
                myfile = data['file']
                fs = FileSystemStorage()
                if fs.exists(myfile.name):
                    return Response({'error' : 'File name already exists!'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                try:
                    form.save()
                    return Response({'data' : form.data}, status=status.HTTP_200_OK)
                except Exception as e:
                     print(e)
                     return Response({'errors' : "erro"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            else:
                return Response("no valid 2", status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            return Response({'errors' : e}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    else:
        return Response({'errors' : form.errors}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@swagger_auto_schema(       
    method='patch',
    operation_description="Edit template variables",
    request_body=openapi.Schema(
        type=openapi.TYPE_OBJECT,
        properties= {'variable_old': openapi.Schema(title="variable_new",type=openapi.TYPE_STRING)},
    ),
)
@api_view(['PATCH'])
@parser_classes([JSONParser])
def edit_variables(request, *args, **kwargs):
    template = TemplateLogic.objects.filter(id=kwargs.get('pk')).values().first()
    editDoc = editDocument(template, request.data, True)
    if template is None:
        return Response({"error" : "Template does not exist!"}, status = status.HTTP_404_NOT_FOUND)
    return Response({'not_found_keys' : editDoc.get('not_found_keys'), 'edited_keys' : editDoc.get('edited_keys')}, status=status.HTTP_200_OK)

def editDocument(template, data, saveDocument):
    doc_variables = get_doc_variables(template)
    if len(doc_variables['vars']) ==  0:
        return Response({"error" : "No variables found"}, status = status.HTTP_404_NOT_FOUND)
    not_found_keys = []
    edited_keys_array = []
    edited_keys = {}
    char_remov = ["-"," "]
    for key_var,var_key_text in enumerate(data):
        if  var_key_text in doc_variables['vars']:
            for key , paragraph in enumerate(doc_variables['doc'].paragraphs):
                if '{{ ' + var_key_text + ' }}' in paragraph.text or '{{' + var_key_text + '}}' in paragraph.text:
                    edited_value = data[var_key_text]
                    for char in char_remov:
                        if saveDocument:
                            edited_keys[var_key_text] = '{{ ' + edited_value + ' }}'
                        else:
                            if type(edited_value) is str:
                                edited_value = edited_value.replace(char, "_")
                                edited_keys[var_key_text] = edited_value
                            if type(edited_value) is list:
                                edited_str = ""
                                for edited in edited_value:
                                    if edited_str == "":
                                        edited_str = edited
                                    else:
                                        edited_str += ", " + edited

                                edited_keys[var_key_text] = edited_str
                    edited_keys_array.append(var_key_text)
        else : 
            not_found_keys.append(var_key_text)

    not_used_vars = []
    for doc_var in doc_variables['vars']:
        if doc_var not in edited_keys_array:
            not_used_vars.append(doc_var)
            edited_keys[doc_var] = '{{ ' + doc_var + ' }}'
    if len(edited_keys_array) > 0 :
        doc_variables['doc'].render(context=edited_keys)
        if saveDocument:
            doc_variables['doc'].save('media/' + template.get('file'))
    return {'not_found_keys' : not_found_keys, 'edited_keys' : edited_keys_array, "doc" : doc_variables['doc']}


@swagger_auto_schema(
    method='get',
    operation_description="Download a Template with all replaced variables from a registered death"
) 
@api_view(['GET'])
def download(request, *args, **kwargs):
    if not TemplateLogic.objects.filter(id=kwargs.get('template_pk')).exists():
        return Response({"error" : "Template does not exist!"},status=status.HTTP_404_NOT_FOUND)
    template_validations = TemplateLogic.objects.filter(id=kwargs.get('template_pk')).values('validations')
    template = TemplateLogic.objects.filter(id=kwargs.get('template_pk')).values().first()
    record = Record.objects.filter(id=kwargs.get('record_pk')).values().first()
   
    if record is None:
        return Response({"error" : "Template does not exist!"}, status = status.HTTP_404_NOT_FOUND)
    if request.data.get('to_send_option') not in dict(SEND_TYPE_CHOICES):
        return Response({"error" : "Field to_send_option needs to be a valid choice", "choices" : dict(SEND_TYPE_CHOICES)}, status = status.HTTP_404_NOT_FOUND)
    
    if request.data.get('to_send_option') == "EMAIL" or request.data.get('to_send_option') == "DOCUMENT_EMAIL":
        from django.core.mail import send_mail, EmailMessage, EmailMultiAlternatives
        from django.template.loader import render_to_string
        from django.utils.html import strip_tags
        context = {'variable1': 'Value 1', 'variable2': 'Value 2'}
        # Get the rendered HTML of the email template
        html_message = render_to_string('email_template.html', context)
        print(html_message)
        # Convert the HTML to plain text for the email body
        plain_message = strip_tags(html_message)
        print(plain_message)
        # Set up the email parameters
        from_email = 'from@example.com'
        html_message = html_message
        group = Group.objects.filter(id= template.get('group_id')).values().first()
        subject = group.get('name') + " - " + template.get('title')
        attachments = []
        if template.get('file'):
            attachments.append("media/" + template.get('file'))
        
        # Send the email
        try:
            email = EmailMultiAlternatives(subject, plain_message, from_email, template.get('send_email_to'))
            email.attach_alternative(html_message, 'text/html')
            for attachment in attachments:
                with open(attachment, 'rb') as file:
                    filename = os.path.basename(attachment)
                    email.attach(filename, file.read())
                    
            email.send()
            #send_mail(subject, message, from_email, template.get('send_email_to'), html_message=html_message)
            print('mandou email')
        except Exception as e:
            print(e)
        
            

    if request.data.get('to_send_option') == "DOCUMENT" or request.data.get('to_send_option') == "DOCUMENT_EMAIL":
        validate_data = run_template_validations(list(template_validations), request.data.get('data'), "CHECK_VALIDATIONS")
        if not validate_data.get('valid') :
            return Response({"success" : False,"errors" : validate_data.get('errors')}, status=status.HTTP_400_BAD_REQUEST)
        if validate_data.get('valid'):
            changeVariablesObject = {}
            keys_missing = []
            doc_variables = get_doc_variables(template)
            getDbKeysToDoc(request, template, template_validations, changeVariablesObject, doc_variables, keys_missing, record)
            if len(keys_missing) > 0:
                return Response({'sucess' : False, "keys_missing" :keys_missing }, status=status.HTTP_400_BAD_REQUEST)
            else:
                editDocRes = editDocument(template, changeVariablesObject, False)  
                document = editDocRes.get('doc')
                buffer = io.BytesIO()
                document.save(buffer)
                buffer.seek(0)
                response = HttpResponse(buffer.read(),content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                import random
                random_number = random.randint(1, 100000000)
                response['Content-Disposition'] = 'attachment; filename='+template.get('title') + '_' + request.user.username + '_' + str(random_number) + '.docx' 
                return response
    return Response({'errors' : "Not Found"}, status=status.HTTP_404_NOT_FOUND)



def getDbKeysToDoc(request, template, template_validations, changeVariablesObject, doc_variables, keys_missing, record):
    for variable in doc_variables['vars']:
            if variable in request.data.get('data'):
                changeVariablesObject[variable] = request.data.get('data')[variable]
    for key,validation in template_validations[0].get('validations').items():
        if not key in changeVariablesObject:
            if "is_field_custom" in validation and not validation.get("is_field_custom") and "db_collection" in validation:
                if validation.get("db_collection") == "USERS":
                    changeVariablesObject[key] = getattr(request.user, validation.get("db_field_reference"), "")
                if validation.get("db_collection") == "GROUPS":
                    group = Group.objects.filter(id= template.get('group_id')).values(validation.get("db_field_reference")).first()
                    if validation.get("db_field_reference") in group:
                        changeVariablesObject[key] = group.get(validation.get("db_field_reference"))
                if validation.get("db_collection") == "RECORDS":
                    if validation.get("db_field_reference") in record:
                        changeVariablesObject[key] = record.get(validation.get("db_field_reference"))
        if not key in changeVariablesObject or changeVariablesObject == "" or changeVariablesObject == None:
            if validation.get("optional"):
                changeVariablesObject[key] = ""
            else:
                keys_missing.append(key)

@swagger_auto_schema(
    method='patch',
    request_body=EditUploadSerializer,
    operation_description="Edit a template"
)    
@api_view(['PATCH'])
@parser_classes([MultiPartParser])
def edit(request, *args, **kwargs):
    template = TemplateLogic.objects.filter(pk=kwargs.get('pk')).first()  
    dup = hasDuplicates(request.POST.get('validations'))
    if dup:
        return dup
    if not request.POST and not request.FILES:
        return Response({"error" : "no data sent"}, status = status.HTTP_400_BAD_REQUEST)
    data = {}
    if request.POST.get('title'):
        data['title'] = request.POST.get('title')
    if request.POST.get('group'):
        data['group'] = request.POST.get('group')
    if request.FILES.get('file'):
        data['file'] = request.FILES.get('file')
    if request.POST.get('validations'):
        data['validations'] = json.loads(request.POST.get('validations'))
    if request.POST.get('send_type'):
        data['send_type'] = json.loads(request.POST.get('send_type'))
    if request.POST.get('send_email_to'):
        data['send_email_to'] = json.loads(request.POST.get('send_email_to'))
    if request.POST.get('send_email_to_cc'):
        data['send_email_to_cc'] = json.loads(request.POST.get('send_email_to_cc'))
    if request.POST.get('send_email_to_bcc'):
        data['send_email_to_bcc'] = json.loads(request.POST.get('send_email_to_bcc'))
    serializer = EditUploadSerializer(data = data,instance=template, partial=True)   
    if serializer.is_valid():
        if template is None:
                return Response({"error" : "Template does not exist!"}, status = status.HTTP_404_NOT_FOUND)
        try:            
            serializer.update(instance = template, validated_data=serializer.validated_data)
            return Response(serializer.data, status = status.HTTP_200_OK)
        except:
            return Response({'error' : "something went wrong updating template","data" : serializer.errors}, status = status.HTTP_404_NOT_FOUND)
    else:
        return Response({'error' : serializer.errors}, status = status.HTTP_400_BAD_REQUEST)   
    

@swagger_auto_schema(
    method='delete',
    operation_description="Delete a Template"
) 
@api_view(['DELETE'])
def remove(request, *args, **kwargs):   
    template = TemplateLogic.objects.filter(id=kwargs.get('pk')).first()
    if template is None:
        return Response({"error" : "Template does not exist!"},status=status.HTTP_404_NOT_FOUND)
    try:
        template.delete()
        return Response({"success" : "Template deleted successfully!"}, status=status.HTTP_204_NO_CONTENT)
    except:
        return Response({"success" : "An error as occured. Try again later!"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
def get_doc_variables(template):
    doc = DocxTemplate('media/' + template.get('file'))
    variables = doc.get_undeclared_template_variables() if doc.get_undeclared_template_variables() else []
    return {
        'vars' : variables,
        'doc' : doc
    }

@swagger_auto_schema(
    method='get',
    operation_description="Get Template Validations"
) 
@api_view(['GET'])
def get_validations(request, *args, **kwargs):
    template_validations = TemplateLogic.objects.filter(id=kwargs.get('pk')).values('validations')
    if template_validations is None:
        return Response({"error" : "Template does not exist!"},status=status.HTTP_404_NOT_FOUND)
    try:
        return Response({"success" : "Template  validations checked successfully!","data" : template_validations}, status=status.HTTP_200_OK)
    except:
        return Response({"success" : "An error as occured. Try again later!"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@swagger_auto_schema(
    method='post',
    operation_description="Check Template Validations"
) 
@api_view(['POST'])
def check_validations(request, *args, **kwargs):
    if not TemplateLogic.objects.filter(id=kwargs.get('pk')).exists():
        return Response({"error" : "Template does not exist!"},status=status.HTTP_404_NOT_FOUND)
    template_validations = TemplateLogic.objects.filter(id=kwargs.get('pk')).values('validations')
    validate_data = run_template_validations(list(template_validations), request.data, "CHECK_VALIDATIONS")
    if not validate_data.get('valid') :
        return Response({"success" : False,"errors" : validate_data.get('errors')}, status=status.HTTP_400_BAD_REQUEST)
    return Response({"success" : True}, status=status.HTTP_200_OK)
    
@swagger_auto_schema(
    method='get',
    operation_description="Get Template"
) 
@api_view(['GET'])
def get_template(request, *args, **kwargs):
    template = TemplateLogic.objects.filter(id=kwargs.get('pk')).values().first()
    if template is None:
        return Response({"error" : "Template does not exist!"},status=status.HTTP_404_NOT_FOUND)
    try:
        return Response({"data" : template}, status=status.HTTP_200_OK)
    except:
        return Response({"error" : "An error as occured. Try again later!"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@swagger_auto_schema(
    method='get',
    operation_description="List group templates"
) 
@api_view(['GET'])
def list_group_templates(request):
    #user = Token.objects.get(key=request.auth.key).user
    if request.user is not None:
        user = User.objects.get(id=request.user.id)
        if (user.is_superuser or user.is_staff) and request.GET.get('group_id'):
            group_id = request.GET.get('group_id')
        elif (user.is_superuser or user.is_staff) and (not request.GET.get('group_id') or request.GET.get('group_id') == ""):
            return Response({'error' : 'group_id field is required'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            group_id = user.group_user_id
        if Group.objects.filter(id=group_id).exists():
            if group_id is not None:
                templates = TemplateLogic.objects.filter(group_id = group_id).values('id', 'title','send_type','file','validations', 'send_type', 'send_email_to', 'send_email_to_cc', 'send_email_to_bcc')
                if templates is not None:
                    if(templates.count() > 0):
                        return Response({'data' : templates}, status=status.HTTP_200_OK)
                    else:
                        return Response({'error' : 'No templates available'}, status=status.HTTP_400_BAD_REQUEST)
        else: 
            if user.is_superuser or user.is_staff:
                return Response({'error' : 'Invalid group_id'}, status=status.HTTP_400_BAD_REQUEST)
            else:
                return Response({'error' : 'User is not assigned to group'}, status=status.HTTP_400_BAD_REQUEST)
    return Response({'error' : 'Invalid user'}, status=status.HTTP_400_BAD_REQUEST)

def hasDuplicates(validations):
    decoder = json.JSONDecoder(object_pairs_hook=parse_object_pairs)
    if validations:
        obj = decoder.decode(validations)
        keys = []
        seen = set()
        for key, value in obj:
            if key not in seen:
                keys.append(key)
                seen.add(key)
            else:
                return Response({"errors" : {'validations' : "Duplicate keys in validations field", 'keys' : keys}}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    return False

def parse_object_pairs(pairs):
    return pairs

@api_view(['GET'])
def sendTestMail(request):
    from django.core.mail import send_mail
    try:
        send_mail(
            'Subject here',
            'Here is the message.',
            'from@example.com',
            ['to@example.com'],
            fail_silently=False,
        )
        return Response({'success' : "Email sent"}, status=status.HTTP_200_OK)
    except Exception as e:
        print(e)
        return Response({'error' : "error sendig email"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    
    
    
    
